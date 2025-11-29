"""
Management command to import Full Stack Java course from Word document
Run with: python manage.py import_java_course
"""
from django.core.management.base import BaseCommand
from django.conf import settings
import os
from pathlib import Path
from learning.models import Course, Module

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class Command(BaseCommand):
    help = 'Imports Full Stack Java course from Word document in static folder'

    def handle(self, *args, **options):
        if not DOCX_AVAILABLE:
            self.stdout.write(
                self.style.ERROR('python-docx is not installed. Please run: pip install python-docx')
            )
            return

        # Find the document in static folder
        static_path = Path(settings.BASE_DIR) / 'static'
        doc_path = static_path / 'Techietact-Course contentv1.0.docx'
        
        if not doc_path.exists():
            self.stdout.write(
                self.style.ERROR(f'Document not found at: {doc_path}')
            )
            return

        self.stdout.write(f'Reading document: {doc_path}')
        
        try:
            doc = Document(str(doc_path))
            content = self.extract_content(doc)
            
            # Create the course
            course, created = Course.objects.get_or_create(
                title='Full Stack Java',
                defaults={
                    'description': 'Complete Full Stack Java Development course covering Java fundamentals, Spring Boot, frontend technologies, and deployment.',
                    'category': 'programming',
                    'is_featured': True,
                }
            )
            
            if created:
                self.stdout.write(self.style.SUCCESS(f'Created course: {course.title}'))
            else:
                self.stdout.write(self.style.WARNING(f'Course already exists: {course.title}'))
                # Delete existing modules to recreate them
                course.modules.all().delete()
                self.stdout.write(self.style.WARNING('Deleted existing modules to recreate from document'))
            
            # Create modules from content
            modules_created = self.create_modules(course, content)
            
            self.stdout.write(
                self.style.SUCCESS(f'\nSuccessfully imported {modules_created} modules for Full Stack Java course!')
            )
            
        except Exception as e:
            self.stdout.write(
                self.style.ERROR(f'Error reading document: {str(e)}')
            )
            import traceback
            traceback.print_exc()

    def extract_content(self, doc):
        """Extract structured content from Word document"""
        content = {
            'modules': []
        }
        
        current_module = None
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if it's a heading (module or section)
            style = para.style.name if para.style else ''
            
            # Module titles are usually Heading 1 or bold/large text
            if 'Heading 1' in style or (para.runs and para.runs[0].bold and len(text) < 100):
                if current_module:
                    content['modules'].append(current_module)
                current_module = {
                    'title': text,
                    'summary': '',
                    'learning_objectives': [],
                    'topics': []
                }
            # Section headings (Heading 2 or similar)
            elif 'Heading 2' in style or (para.runs and para.runs[0].bold):
                current_section = text
                if current_module:
                    current_module['topics'].append(text)
            # Regular content
            else:
                if current_module:
                    if current_section and 'objective' in current_section.lower():
                        current_module['learning_objectives'].append(text)
                    elif current_section and ('topic' in current_section.lower() or 'content' in current_section.lower()):
                        current_module['topics'].append(text)
                    elif not current_module['summary']:
                        current_module['summary'] = text
                    else:
                        # Add to topics as additional content
                        current_module['topics'].append(text)
        
        # Add last module
        if current_module:
            content['modules'].append(current_module)
        
        # If no structured content found, create modules from paragraphs
        if not content['modules']:
            content = self.extract_simple_content(doc)
        
        return content

    def extract_simple_content(self, doc):
        """Extract content when structure is not clear - create modules from major sections"""
        content = {
            'modules': []
        }
        
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # Common Full Stack Java module structure
        module_titles = [
            'Java Fundamentals',
            'Object-Oriented Programming',
            'Java Collections Framework',
            'Database Connectivity (JDBC)',
            'Spring Framework',
            'Spring Boot',
            'RESTful APIs',
            'Frontend Integration',
            'Security & Authentication',
            'Deployment & DevOps'
        ]
        
        current_module_idx = 0
        current_content = []
        
        for para in paragraphs:
            # Check if paragraph matches a module title
            is_module_title = any(title.lower() in para.lower() for title in module_titles)
            
            if is_module_title and current_content:
                # Save previous module
                if current_module_idx < len(module_titles):
                    content['modules'].append({
                        'title': module_titles[current_module_idx],
                        'summary': current_content[0] if current_content else 'Learn essential concepts',
                        'learning_objectives': current_content[1:3] if len(current_content) > 1 else ['Master key concepts', 'Apply knowledge practically'],
                        'topics': current_content[3:8] if len(current_content) > 3 else ['Introduction', 'Core concepts', 'Practice exercises']
                    })
                    current_module_idx += 1
                    current_content = []
            
            if para and len(para) > 10:  # Skip very short lines
                current_content.append(para)
        
        # Add remaining content as last module
        if current_content and current_module_idx < len(module_titles):
            content['modules'].append({
                'title': module_titles[current_module_idx],
                'summary': current_content[0] if current_content else 'Learn essential concepts',
                'learning_objectives': current_content[1:3] if len(current_content) > 1 else ['Master key concepts', 'Apply knowledge practically'],
                'topics': current_content[3:8] if len(current_content) > 3 else ['Introduction', 'Core concepts', 'Practice exercises']
            })
        
        # If still no modules, create default structure
        if not content['modules']:
            content['modules'] = [
                {
                    'title': title,
                    'summary': f'Learn {title.lower()} concepts and best practices',
                    'learning_objectives': ['Understand core concepts', 'Apply knowledge in projects', 'Build practical skills'],
                    'topics': ['Introduction', 'Core concepts', 'Advanced topics', 'Best practices', 'Hands-on exercises']
                }
                for title in module_titles[:10]  # First 10 modules
            ]
        
        return content

    def create_modules(self, course, content):
        """Create modules from extracted content"""
        modules_created = 0
        
        for order, module_data in enumerate(content['modules'], start=1):
            learning_objectives = '\n'.join(module_data.get('learning_objectives', []))
            topics = '\n'.join(module_data.get('topics', []))
            
            module = Module.objects.create(
                course=course,
                order=order,
                title=module_data.get('title', f'Module {order}'),
                summary=module_data.get('summary', 'Learn essential concepts'),
                learning_objectives=learning_objectives if learning_objectives else 'Master key concepts\nApply knowledge practically\nBuild real-world projects',
                topics=topics if topics else 'Introduction\nCore concepts\nAdvanced topics\nPractice exercises'
            )
            modules_created += 1
            self.stdout.write(
                self.style.SUCCESS(f'  Created module {order}: {module.title}')
            )
        
        return modules_created

